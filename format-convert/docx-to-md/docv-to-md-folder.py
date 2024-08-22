import os
import PyPDF2
import docx
import re
from docx.enum.style import WD_STYLE_TYPE

# Define your input and output paths here
INPUT_FOLDER = r"/Users/kadirkatirci/Desktop/sil/docx-to-md/docx"
OUTPUT_FOLDER = r"/Users/kadirkatirci/Desktop/sil/docx-to-md/md"

def pdf_to_markdown(pdf_path):
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n\n"
    
    # Basic formatting
    text = re.sub(r'^([A-Z\s]+)$', r'# \1', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*•\s*', '* ', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*(\d+)\.\s*', r'\1. ', text, flags=re.MULTILINE)
    
    return text

def docx_to_markdown(docx_path):
    doc = docx.Document(docx_path)
    markdown_text = ""

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name[-1])
            markdown_text += f"{'#' * level} {text}\n\n"
        elif paragraph.style.name == 'List Paragraph':
            if re.match(r'^\d+\.', text):
                markdown_text += f"{text}\n"
            else:
                markdown_text += f"* {text}\n"
        else:
            markdown_text += f"{text}\n\n"

    for table in doc.tables:
        markdown_text += "| "
        markdown_text += " | ".join([cell.text for cell in table.rows[0].cells])
        markdown_text += " |\n"
        markdown_text += "| " + " | ".join(["---" for _ in table.rows[0].cells]) + " |\n"
        for row in table.rows[1:]:
            markdown_text += "| "
            markdown_text += " | ".join([cell.text for cell in row.cells])
            markdown_text += " |\n"
        markdown_text += "\n"

    return markdown_text

def convert_files(input_folder, output_folder):
    total_files = 0
    converted_files = 0
    error_files = []

    # Count total files
    for root, dirs, files in os.walk(input_folder):
        for file in files:
            if file.endswith(('.pdf', '.docx')):
                total_files += 1

    print(f"Found {total_files} files to convert.")

    # Convert files
    for root, dirs, files in os.walk(input_folder):
        for file in files:
            if file.endswith(('.pdf', '.docx')):
                input_path = os.path.join(root, file)
                relative_path = os.path.relpath(input_path, input_folder)
                output_path = os.path.join(output_folder, os.path.splitext(relative_path)[0] + '.md')

                # Create output directory if it doesn't exist
                os.makedirs(os.path.dirname(output_path), exist_ok=True)

                try:
                    if file.endswith('.pdf'):
                        markdown_content = pdf_to_markdown(input_path)
                    else:  # .docx
                        markdown_content = docx_to_markdown(input_path)

                    with open(output_path, 'w', encoding='utf-8') as md_file:
                        md_file.write(markdown_content)

                    converted_files += 1
                    print(f"Converted ({converted_files}/{total_files}): {input_path} to {output_path}")
                except PyPDF2.errors.PdfReadError as e:
                    error_message = f"PDF Read Error for {input_path}: {str(e)}"
                    print(error_message)
                    error_files.append((input_path, error_message))
                except docx.opc.exceptions.PackageNotFoundError as e:
                    error_message = f"DOCX Read Error for {input_path}: {str(e)}"
                    print(error_message)
                    error_files.append((input_path, error_message))
                except PermissionError as e:
                    error_message = f"Permission Error for {input_path}: {str(e)}"
                    print(error_message)
                    error_files.append((input_path, error_message))
                except Exception as e:
                    error_message = f"Unexpected Error converting {input_path}: {str(e)}"
                    print(error_message)
                    error_files.append((input_path, error_message))

    print(f"\nConversion complete. Successfully converted {converted_files} out of {total_files} files.")
    
    if error_files:
        print(f"\nEncountered errors with {len(error_files)} files:")
        for file, error in error_files:
            print(f"- {file}: {error}")

if __name__ == "__main__":
    # Validate input folder
    if not os.path.isdir(INPUT_FOLDER):
        print(f"Error: Input folder '{INPUT_FOLDER}' does not exist or is not a directory.")
        exit(1)

    # Create output folder if it doesn't exist
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    convert_files(INPUT_FOLDER, OUTPUT_FOLDER)