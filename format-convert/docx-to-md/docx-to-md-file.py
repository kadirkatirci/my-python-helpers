import docx
from docx.enum.style import WD_STYLE_TYPE
import re

def docx_to_markdown(docx_path, output_path):
    # Open the docx file
    doc = docx.Document(docx_path)
    
    markdown_text = ""

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        # Check paragraph style
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name[-1])
            markdown_text += f"{'#' * level} {text}\n\n"
        elif paragraph.style.name == 'List Paragraph':
            # Check if it's a numbered list
            if re.match(r'^\d+\.', text):
                markdown_text += f"{text}\n"
            else:
                markdown_text += f"* {text}\n"
        else:
            markdown_text += f"{text}\n\n"

    # Handle tables
    for table in doc.tables:
        markdown_text += "| "
        markdown_text += " | ".join([cell.text for cell in table.rows[0].cells])
        markdown_text += " |\n"
        markdown_text += "| " + " | ".join(["---" for _ in table.rows[0].cells]) + " |\n"
        for row in table.rows[1:]:
            markdown_text += "| "
            markdown_text += " | ".join([cell.text for cell in row.cells])
            markdown_text += " |\n"
        markdown_text += "\n"

    # Write to markdown file
    with open(output_path, 'w', encoding='utf-8') as md_file:
        md_file.write(markdown_text)

# Usage
docx_path = 'AB.docx'
output_path = 'AB.md'
docx_to_markdown(docx_path, output_path)
print(f"Conversion complete. Markdown file saved as {output_path}")